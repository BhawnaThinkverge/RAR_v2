import pandas as pd
import streamlit as st
import os
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain.callbacks import get_openai_callback
from langgraph.graph import StateGraph, END
from typing import TypedDict, List, Dict
import json
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

load_dotenv()
os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")

# ==================== SECTOR DATA PATHS =======================================================================================================
SECTOR_PATHS = {
    "Manufacturing": r"C:\Users\JaswanthPalla\Desktop\VS  Code\ProbBid\Codebase\tester_1_20-08-2025\RAR-V1\Manufacturing_Sector 1.xlsx",
    "Cyber Security": None,
    "Others": None
}

# ==================== PERSONA CONFIGURATION ==================================================================================================
PERSONA_PROMPTS = {
    "Board / Executive Sponsor": {
        "icon": "🏛️",
        "tone": "Strategic, financial, risk-focused",
        "focus": "ROI, auditability, large spends, strategic alignment, business case"
    },
    "CIO / CTO": {
        "icon": "🖥️",
        "tone": "Technical strategy & architecture",
        "focus": "Integration, TCO, standards, scalability, IT roadmap"
    },
    "CDO / Head of Digital / Innovation": {
        "icon": "🚀",
        "tone": "Innovation & transformation",
        "focus": "AI-readiness, UX, modernization, data access, speed to value"
    },
    "CISO / Head of Cyber": {
        "icon": "🔒",
        "tone": "Security & compliance first",
        "focus": "Risk exposure, data protection, compliance, vendor security posture"
    },
    "Enterprise / Solution Architect": {
        "icon": "⚙️",
        "tone": "Deep technical architecture",
        "focus": "Interoperability, APIs, data models, technical debt, migration paths"
    }
}

# ==================== STATE DEFINITION ======================================================================================================
class ComparisonState(TypedDict):
    input_tool: Dict
    compare_tools: List[Dict]
    persona: str
    report: str
    estimated_tokens: int
    actual_tokens: Dict[str, int]
    estimation_done: bool

# ==================== TOKEN ESTIMATION ======================================================================================================
def prepare_data(state: ComparisonState) -> ComparisonState:
    if state.get("estimation_done", False):
        return state
    
    selected_json = json.dumps(state["input_tool"], separators=(',', ':'))
    data_json = json.dumps(state["compare_tools"], separators=(',', ':'))
    
    # Rough but better estimation
    base_prompt = f"""Compare {selected_json} vs {len(state["compare_tools"])} tools as {state.get('persona', 'Executive')} 
    with SWOT, GAP, Market Share, and Recommendation."""
    
    total_chars = len(base_prompt) + len(selected_json) + len(data_json) + 3000  # buffer for full prompt
    state["estimated_tokens"] = int(total_chars / 4)
    state["estimation_done"] = True
    return state

# ==================== PERSONA-TAILORED PROMPT WITH ORIGINAL STRUCTURE =======================================================================
def generate_comparison(state: ComparisonState) -> ComparisonState:
    selected_persona = state.get("persona", "Board / Executive Sponsor")

    # Persona-specific system instructions
    system_prompts = {
        "Board / Executive Sponsor": """You are a senior strategy consultant briefing the Board and CEO. 
        Use crisp, financial, risk-aware language. Focus on ROI, strategic alignment, business justification, and auditability. 
        Emphasize financial impact, market positioning, and executive decision implications. Avoid deep technical jargon.""",

        "CIO / CTO": """You are advising the CIO/CTO on enterprise architecture and IT strategy. 
        Focus on integration complexity, total cost of ownership, scalability, standards compliance, and long-term operating impact.""",

        "CDO / Head of Digital / Innovation": """You are the Chief Digital Officer driving transformation. 
        Emphasize AI/ML readiness, user adoption, speed of innovation, data democratization, and modernization potential.""",

        "CISO / Head of Cyber": """You are the CISO. Prioritize security posture, compliance risks, data sovereignty, 
        identity management, and vendor risk exposure. Be direct about red flags and material threats.""",

        "Enterprise / Solution Architect": """You are a principal enterprise architect. Provide deep technical analysis on 
        integration patterns, API strategy, data modeling, deployment architecture, technical debt, and migration feasibility."""
    }

    # MAIN PROMPT — Keeps your ORIGINAL desired structure + persona tone
    prompt_template = """You are an expert enterprise software analyst generating a comparison report for the selected tool 
    "{tool_name}" ({tool_name} by {vendor}) versus top competing PLM/CAD tools in the market.

    Audience: {persona}
    Tone & Focus: {tone}. Prioritize: {focus}.

    Use only verifiable data from vendor sites, Gartner, Forrester, case studies, and public benchmarks.

    Internally score all tools (including selected) on a 1–5 scale across:
    - Market Presence (10%)
    - Cost (20%)
    - Integration (25%)
    - Features (30%)
    - Efficiency (15%)

    Calculate weighted scores and rank all tools. Identify the true top 5. DO NOT reveal scores, rankings, or tool names except the selected one.

    Output EXACTLY these sections in this order, in clean markdown:

    **Market Share Overview**
    - 3–5 bullets comparing current global/regional share, growth rates (CAGR), and forecasts (2024–2028) for the selected tool vs. top 5 leaders collectively
    - Use real numbers where possible (e.g., "holds ~6% global share vs. leaders at 18–28%, growing at 9% CAGR vs. market 13%")

    **SWOT Analysis** (for the selected tool relative to top 5)
    | **Strengths** | **Weaknesses** |
    |---------------|----------------|
    | - Bullet 1
    | - Bullet 2
    | - Bullet 3 | - Bullet 1
    |               | - Bullet 2
    |               | - Bullet 3 |

    | **Opportunities** | **Threats** |
    |-------------------|-----------|
    | - Bullet 1
    | - Bullet 2
    | - Bullet 3 | - Bullet 1
    |                   | - Bullet 2
    |                   | - Bullet 3 |

    **GAP Analysis**
    - 4–6 consultant-style bullets highlighting critical feature/innovation gaps vs. top 5
    - Each bullet: bold gap name + data-backed explanation + action recommendation
    - Example: "- **Generative Design Gap**: Lacks native AI modeling (Autodesk Fusion strength); invest in R&D or partner to close 18-month lag."

    **Recommendation** (tailored to {persona})
    - First bullet: Name the #1 ranked tool (or selected if it wins)
    - Next bullets: Key strengths/weaknesses of selected vs. best, with data justification
    - Final bullet: Clear **Switch: Yes/No/Maybe (Pilot)** + 1-sentence reasoning tied to role priorities

    Write as if presenting directly to the {persona}. Use their language and concerns. Be decisive and action-oriented."""

    chat_llm = ChatOpenAI(model="gpt-4o", temperature=0.2, max_tokens=2000)

    prompt = ChatPromptTemplate.from_messages([
        ("system", system_prompts.get(selected_persona, system_prompts["Board / Executive Sponsor"])),
        ("human", prompt_template)
    ])

    chain = (
        {
            "tool_name": lambda x: x["input_tool"].get("Tool Name", "Selected Tool"),
            "vendor": lambda x: x["input_tool"].get("Vendor", "Unknown"),
            "persona": lambda x: x.get("persona", "Board / Executive Sponsor"),
            "tone": lambda x: PERSONA_PROMPTS[x.get("persona", "Board / Executive Sponsor")]["tone"],
            "focus": lambda x: PERSONA_PROMPTS[x.get("persona", "Board / Executive Sponsor")]["focus"]
        }
        | prompt
        | chat_llm
        | StrOutputParser()
    )

    with get_openai_callback() as cb:
        response = chain.invoke(state)

    state["actual_tokens"] = {
        "prompt_tokens": cb.prompt_tokens,
        "completion_tokens": cb.completion_tokens,
        "total_tokens": cb.total_tokens
    }
    state["report"] = response
    return state

# ==================== DOCX EXPORT (Improved Table Handling) =========================================================================
def create_docx_from_markdown(md_content):
    doc = Document()
    doc.add_heading('Research & Analysis Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    lines = md_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            doc.add_paragraph()
            i += 1
            continue

        if line.startswith("**") and line.endswith("**"):
            doc.add_heading(line[2:-2], level=1)
        elif line.startswith("| **Strengths**") or line.startswith("| **Opportunities**"):
            # Start of SWOT table
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1

            if len(table_lines) >= 6:
                doc.add_paragraph()
                table = doc.add_table(rows=6, cols=2)
                table.style = 'Table Grid'

                # Fill cells with preserved line breaks
                cells_content = [
                    [tl.split("|")[1].strip(), tl.split("|")[2].strip()] if len(tl.split("|")) > 2 else ["", ""]
                    for tl in table_lines[:6]
                ]

                for row_idx, row in enumerate(cells_content):
                    for col_idx, cell_text in enumerate(row):
                        cell = table.cell(row_idx, col_idx)
                        cell.text = ""
                        for bullet in cell_text.splitlines():
                            if bullet.strip():
                                p = cell.add_paragraph(bullet.strip(), style='List Bullet' if row_idx > 1 else None)
                                if row_idx <= 1:
                                    p.runs[0].bold = True
        elif line.startswith("- ") or line.startswith("• "):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif "**" in line:
            p = doc.add_paragraph()
            parts = line.split("**")
            for j, part in enumerate(parts):
                run = p.add_run(part)
                if j % 2 == 1:
                    run.bold = True
        else:
            doc.add_paragraph(line)
        i += 1

    doc.add_page_break()
    footer = doc.add_paragraph(f'Generated on {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M")}')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].italic = True

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

# ==================== GRAPH & CACHING =================================================================================================
def build_graph():
    wf = StateGraph(ComparisonState)
    wf.add_node("prepare", prepare_data)
    wf.add_node("compare", generate_comparison)
    wf.set_entry_point("prepare")
    wf.add_edge("prepare", "compare")
    wf.add_edge("compare", END)
    return wf.compile()

graph = build_graph()

@st.cache_data
def load_sheet_names(_excel_path):
    xls = pd.ExcelFile(_excel_path)
    return xls.sheet_names

def load_tools(_excel_path, _sheet_name):
    df = pd.read_excel(_excel_path, sheet_name=_sheet_name)
    df.columns = df.columns.str.strip()
    return df

# ==================== STREAMLIT UI =====================================================================================================
st.set_page_config(layout="wide", page_title="RAR - Research & Analysis", page_icon="🔍")
st.title("🔍 Research & Analysis Report Generator")

for key in ['selected_sector', 'selected_sheet', 'input_choice', 'selected_persona', 'comparison_result']:
    if key not in st.session_state:
        st.session_state[key] = None

if st.session_state.selected_sector is None:
    st.session_state.selected_sector = "Manufacturing"
if st.session_state.selected_persona is None:
    st.session_state.selected_persona = "Board / Executive Sponsor"

col1, col2 = st.columns([1, 2])

with col1:
    st.subheader("⚙️ Configuration")

    selected_sector = st.selectbox("Sector:", list(SECTOR_PATHS.keys()), 
                                   index=list(SECTOR_PATHS.keys()).index(st.session_state.selected_sector))

    if selected_sector != st.session_state.selected_sector:
        st.session_state.selected_sector = selected_sector
        st.session_state.selected_sheet = None
        st.session_state.input_choice = None
        st.session_state.comparison_result = None
        st.rerun()

    excel_path = SECTOR_PATHS.get(selected_sector)

    if excel_path is None:
        st.warning("Data in progress for this sector.")
    else:
        sheet_names = load_sheet_names(excel_path)
        if st.session_state.selected_sheet is None:
            st.session_state.selected_sheet = sheet_names[0]

        selected_sheet = st.selectbox("Stage / Sheet:", sheet_names,
                                      index=sheet_names.index(st.session_state.selected_sheet))

        if selected_sheet != st.session_state.selected_sheet:
            st.session_state.selected_sheet = selected_sheet
            st.session_state.input_choice = None
            st.session_state.comparison_result = None
            st.rerun()

        df = load_tools(excel_path, selected_sheet)
        tools = df['Tool Name'].dropna().unique().tolist()

        if len(tools) > 0:
            if st.session_state.input_choice not in tools:
                st.session_state.input_choice = tools[0]

            input_choice = st.selectbox("Tool to Analyze:", tools,
                                        index=tools.index(st.session_state.input_choice),
                                        key="tool_select")

            if input_choice != st.session_state.input_choice:
                st.session_state.input_choice = input_choice
                st.session_state.comparison_result = None

            st.markdown("---")
            st.markdown("#### 👤 Report Audience")

            persona_options = list(PERSONA_PROMPTS.keys())
            selected_persona = st.selectbox(
                "Tailor report for:",
                options=persona_options,
                format_func=lambda x: f"{PERSONA_PROMPTS[x]['icon']} {x}",
                index=persona_options.index(st.session_state.selected_persona),
                help="Changes tone, depth, and focus of the entire report",
                key="persona_select_final"
            )

            if selected_persona != st.session_state.selected_persona:
                st.session_state.selected_persona = selected_persona
                st.session_state.comparison_result = None

            tool_row = df[df['Tool Name'] == input_choice].iloc[0]
            st.markdown(f"**Selected Tool:** {tool_row['Tool Name']} ({tool_row['Vendor']})")

            compare_button = st.button("🚀 Generate Report", type="primary", use_container_width=True)

            if compare_button:
                with st.spinner(f"Generating report for {selected_persona}..."):
                    input_tool = tool_row[['Tool Name', 'Tool Type', 'Vendor', 'Region', 'Cost',
                                           'Number of Integrations', 'Features', 'Efficiency', 'Source']].to_dict()
                    compare_tools = df[df['Tool Name'] != input_choice][['Tool Name', 'Tool Type', 'Vendor', 'Region', 'Cost',
                                                                         'Number of Integrations', 'Features', 'Efficiency', 'Source']].to_dict('records')

                    initial_state = {
                        "input_tool": input_tool,
                        "compare_tools": compare_tools,
                        "persona": selected_persona,
                        "report": "",
                        "estimated_tokens": 0,
                        "actual_tokens": {},
                        "estimation_done": False
                    }

                    prep_state = prepare_data(initial_state)
                    result = graph.invoke(prep_state)
                    st.session_state.comparison_result = result
                    st.rerun()

with col2:
    st.subheader("📊 Generated Report")

    if st.session_state.comparison_result:
        result = st.session_state.comparison_result
        report = result["report"]

        st.success(f"Report generated for {st.session_state.selected_persona}!")

        docx_bytes = create_docx_from_markdown(report)
        st.download_button(
            "📥 Download as Word (.docx)",
            data=docx_bytes,
            file_name=f"RAR_{st.session_state.input_choice.replace(' ', '_')}_{st.session_state.selected_persona.replace(' ', '_').replace('/', '')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

        st.markdown("---")
        with st.container(border=True):
            st.markdown(report, unsafe_allow_html=True)
    else:
        st.info("Select a tool and persona, then click **Generate Report** to see the output!")

st.sidebar.success("Persona-tailored structured reports ready!")
